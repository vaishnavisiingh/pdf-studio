from fastapi import APIRouter, HTTPException
from fastapi.responses import Response
import api.routes_document as doc_module
from core.idrep import IDRepNodeType
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

router = APIRouter(prefix="/export", tags=["export"])


def idrep_to_docx(idrep) -> bytes:
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Sort nodes by page then sort_order
    text_types = {
        IDRepNodeType.HEADING,
        IDRepNodeType.PARAGRAPH,
        IDRepNodeType.CAPTION,
        IDRepNodeType.REFERENCE,
        IDRepNodeType.EQUATION,
        IDRepNodeType.TABLE,
    }

    nodes = sorted(
        [n for n in idrep._node_index.values() if n.node_type in text_types],
        key=lambda n: (n.page_number, n.sort_order)
    )

    for node in nodes:
        if node.node_type == IDRepNodeType.HEADING:
            level = min(node.heading_level or 1, 4)
            p = doc.add_heading(node.text or "", level=level)

        elif node.node_type == IDRepNodeType.PARAGRAPH:
            p = doc.add_paragraph(node.text or "")
            if node.is_bold:
                for run in p.runs:
                    run.bold = True
            if node.is_italic:
                for run in p.runs:
                    run.italic = True

        elif node.node_type == IDRepNodeType.CAPTION:
            p = doc.add_paragraph(node.text or "")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        elif node.node_type == IDRepNodeType.REFERENCE:
            p = doc.add_paragraph(node.text or "")
            for run in p.runs:
                run.font.size = Pt(9)

        elif node.node_type == IDRepNodeType.EQUATION:
            p = doc.add_paragraph(f"[Equation: {node.text or ''}]")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.italic = True

        elif node.node_type == IDRepNodeType.TABLE:
            rows = sorted(node.children, key=lambda n: n.sort_order)
            if not rows:
                continue
            col_count = max(len(r.children) for r in rows) if rows else 1
            table = doc.add_table(rows=len(rows), cols=col_count)
            table.style = "Table Grid"
            for r_idx, row_node in enumerate(rows):
                cells = sorted(row_node.children, key=lambda n: n.sort_order)
                for c_idx, cell_node in enumerate(cells):
                    if c_idx < col_count:
                        table.cell(r_idx, c_idx).text = cell_node.text or ""

        doc.add_paragraph("")  # spacing

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.get("/{doc_id}/docx")
async def export_docx(doc_id: str):
    session = doc_module._sessions.get(doc_id)
    if not session:
        raise HTTPException(404, "Document not found")

    try:
        docx_bytes = idrep_to_docx(session["idrep"])
    except Exception as e:
        raise HTTPException(500, f"Export failed: {str(e)}")

    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=export.docx"}
    )


@router.get("/{doc_id}/txt")
async def export_txt(doc_id: str):
    session = doc_module._sessions.get(doc_id)
    if not session:
        raise HTTPException(404, "Document not found")

    idrep = session["idrep"]
    lines = []

    nodes = sorted(
        [n for n in idrep._node_index.values()
         if n.node_type in {IDRepNodeType.HEADING, IDRepNodeType.PARAGRAPH,
                            IDRepNodeType.CAPTION, IDRepNodeType.REFERENCE}],
        key=lambda n: (n.page_number, n.sort_order)
    )

    for node in nodes:
        if node.text:
            lines.append(node.text)
            lines.append("")

    return Response(
        content="\n".join(lines).encode("utf-8"),
        media_type="text/plain",
        headers={"Content-Disposition": "attachment; filename=export.txt"}
    )


@router.get("/{doc_id}/html")
async def export_html(doc_id: str):
    session = doc_module._sessions.get(doc_id)
    if not session:
        raise HTTPException(404, "Document not found")

    idrep = session["idrep"]
    text_types = {
        IDRepNodeType.HEADING,
        IDRepNodeType.PARAGRAPH,
        IDRepNodeType.CAPTION,
        IDRepNodeType.REFERENCE,
        IDRepNodeType.TABLE,
    }

    nodes = sorted(
        [n for n in idrep._node_index.values() if n.node_type in text_types],
        key=lambda n: (n.page_number, n.sort_order)
    )

    html_parts = ["""<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<title>Exported Document</title>
<style>
  body { font-family: Georgia, serif; max-width: 800px; margin: 40px auto; padding: 0 20px; line-height: 1.6; color: #111; }
  h1 { font-size: 2em; border-bottom: 2px solid #333; padding-bottom: 8px; }
  h2 { font-size: 1.5em; margin-top: 32px; }
  h3 { font-size: 1.2em; margin-top: 24px; }
  p { margin: 12px 0; text-align: justify; }
  .caption { font-style: italic; color: #666; font-size: 0.9em; text-align: center; }
  .reference { font-size: 0.85em; color: #444; margin: 4px 0; }
  table { border-collapse: collapse; width: 100%; margin: 16px 0; }
  th { background: #1e2a45; color: white; padding: 8px 12px; text-align: left; }
  td { border: 1px solid #dee2e6; padding: 7px 12px; }
  tr:nth-child(even) { background: #f8f9fa; }
</style>
</head>
<body>"""]

    for node in nodes:
        text = (node.text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        if node.node_type == IDRepNodeType.HEADING:
            level = min(node.heading_level or 1, 4)
            html_parts.append(f"<h{level}>{text}</h{level}>")

        elif node.node_type == IDRepNodeType.PARAGRAPH:
            if text.strip():
                html_parts.append(f"<p>{text}</p>")

        elif node.node_type == IDRepNodeType.CAPTION:
            html_parts.append(f'<p class="caption">{text}</p>')

        elif node.node_type == IDRepNodeType.REFERENCE:
            html_parts.append(f'<p class="reference">{text}</p>')

        elif node.node_type == IDRepNodeType.TABLE:
            rows = sorted(node.children, key=lambda n: n.sort_order)
            if rows:
                html_parts.append("<table>")
                for i, row_node in enumerate(rows):
                    cells = sorted(row_node.children, key=lambda n: n.sort_order)
                    tag = "th" if i == 0 else "td"
                    html_parts.append("<tr>")
                    for cell in cells:
                        cell_text = (cell.text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                        html_parts.append(f"<{tag}>{cell_text}</{tag}>")
                    html_parts.append("</tr>")
                html_parts.append("</table>")

    html_parts.append("</body></html>")

    return Response(
        content="\n".join(html_parts).encode("utf-8"),
        media_type="text/html",
        headers={"Content-Disposition": "attachment; filename=export.html"}
    )
